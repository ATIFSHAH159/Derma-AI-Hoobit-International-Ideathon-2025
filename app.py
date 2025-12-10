# app.py - Flask Backend for Skin Cancer Detection
from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
from dotenv import load_dotenv
load_dotenv()

import tensorflow as tf
from tensorflow import keras
from tensorflow.keras.applications.efficientnet import preprocess_input
import numpy as np
from PIL import Image
import io
import os
import time
import re
import csv
import pandas as pd
import google.generativeai as genai

# Optional report generation libs
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

app = Flask(__name__)
CORS(app)  # Enable CORS for cross-origin requests

# Configuration
MODEL_PATH = 'skin_cancer_final_model.keras'
IMG_SIZE = (224, 224)
MODEL_ACCURACY = 85.0

# Gemini API Configuration
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY', '')  # Set via environment variable or .env file
GEMINI_MODEL = 'gemini-2.0-flash-001'  # Using Gemini 2.0 Flash (stable version)

# Class names (MUST match the order from your training - categorical order)
CLASS_NAMES = [
    'Actinic keratoses',
    'Basal cell carcinoma',
    'Benign keratosis-like lesions',
    'Dermatofibroma',
    'Melanocytic nevi',
    'Melanoma',
    'Vascular lesions'
]

# Mapping from class names to CSV condition codes
CONDITION_CODE_MAP = {
    'Actinic keratoses': 'akiec',
    'Basal cell carcinoma': 'bcc',
    'Benign keratosis-like lesions': 'bkl',
    'Dermatofibroma': 'df',
    'Melanocytic nevi': 'nv',
    'Melanoma': 'mel',
    'Vascular lesions': 'vasc'
}

# Doctors CSV file path
DOCTORS_CSV_PATH = 'doctors_cancer_pakistan_full.csv'

# Load model
print("="*60)
print("SKIN CANCER DETECTION API - INITIALIZING")
print("="*60)

try:
    model = keras.models.load_model(MODEL_PATH)
    print(f"✓ Model loaded successfully: {MODEL_PATH}")
    print(f"✓ Model input shape: {model.input_shape}")
    print(f"✓ Number of classes: {len(CLASS_NAMES)}")
except Exception as e:
    print(f"❌ Error loading model: {e}")
    model = None

print("="*60)

# Load doctors data
doctors_df = None
try:
    if os.path.exists(DOCTORS_CSV_PATH):
        doctors_df = pd.read_csv(DOCTORS_CSV_PATH)
        print(f"✓ Doctors database loaded: {len(doctors_df)} doctors")
    else:
        print(f"⚠️ Warning: Doctors CSV file not found: {DOCTORS_CSV_PATH}")
except Exception as e:
    print(f"⚠️ Warning: Failed to load doctors database: {e}")
    doctors_df = None

# Initialize Gemini API
if GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        print(f"✓ Gemini API configured successfully")
        print(f"✓ Using model: {GEMINI_MODEL}")
        print(f"  Note: Free tier has rate limits. Retry logic enabled for quota errors.")
    except Exception as e:
        print(f"⚠️ Warning: Gemini API configuration failed: {e}")
else:
    print("⚠️ Warning: GEMINI_API_KEY not set. Medical analysis will be disabled.")
    print("   Set it via environment variable: export GEMINI_API_KEY='your-api-key'")

print("="*60)


def get_medical_analysis(cancer_name, confidence_score, model_accuracy, max_retries=3):
    """
    Get specialized medical LLM analysis from Google Gemini 2.0 Flash
    Includes retry logic for rate limit errors (429)
    
    Args:
        cancer_name: Name of the detected skin lesion/cancer type
        confidence_score: Confidence score from the model (0-1)
        model_accuracy: Overall model accuracy percentage
        max_retries: Maximum number of retry attempts for rate limit errors
    
    Returns:
        Dictionary with structured medical analysis
    """
    if not GEMINI_API_KEY:
        return {
            'error': 'Gemini API key not configured',
            'analysis': None
        }
    
    # Create the prompt for medical analysis (optimized for token usage)
    prompt = f"""You are a specialized medical AI assistant. Provide dermatological analysis.

DETECTION RESULTS:
- Condition: {cancer_name}
- Confidence: {confidence_score:.2%}
- Model Accuracy: {model_accuracy}%

Provide analysis in this format:

1. **Type of Condition**: What {cancer_name} is, nature (benign/malignant/precancerous), characteristics.

2. **Stage Indicators**: Early/late stage indicators or why staging doesn't apply.

3. **Medical Literature Evidence**: Prevalence, demographics, risk factors, clinical presentation, diagnostic criteria.

4. **Treatment Steps**: Medical interventions, surgical options (if applicable), follow-up care.

5. **Home-Care Recommendations**: Safe skin care practices, monitoring guidelines, what to watch for.

6. **Non-Prescription Medicine Guidance**: OTC topical treatments (if safe), pain management, self-medication warnings.

7. **Urgency Indicators - When to See a Doctor**: Immediate/urgent signs, when to schedule appointment, routine follow-up, red flag symptoms.

DISCLAIMER: This is AI-assisted analysis, not a replacement for professional medical diagnosis. Consult a qualified dermatologist.

Format with clear section headers. Be concise but comprehensive."""

    # Initialize the model
    gemini_model = genai.GenerativeModel(GEMINI_MODEL)
    
    # Retry logic for rate limit errors
    for attempt in range(max_retries):
        try:
            print(f"📡 Requesting medical analysis from Gemini for: {cancer_name} (Attempt {attempt + 1}/{max_retries})")
            response = gemini_model.generate_content(prompt)
            
            analysis_text = response.text
            
            print(f"✓ Medical analysis received from Gemini")
            
            return {
                'error': None,
                'analysis': analysis_text,
                'model_used': GEMINI_MODEL
            }
            
        except Exception as e:
            error_str = str(e)
            
            # Check if it's a rate limit error (429)
            if '429' in error_str or 'quota' in error_str.lower() or 'rate limit' in error_str.lower():
                if attempt < max_retries - 1:
                    # Extract retry delay if available, otherwise use exponential backoff
                    retry_delay = 15 * (2 ** attempt)  # 15s, 30s, 60s
                    
                    # Try to extract retry delay from error message
                    if 'retry in' in error_str.lower() or 'retry_delay' in error_str.lower():
                        # Look for seconds in error message
                        seconds_match = re.search(r'(\d+\.?\d*)\s*seconds?', error_str, re.IGNORECASE)
                        if seconds_match:
                            retry_delay = int(float(seconds_match.group(1))) + 2  # Add 2 seconds buffer
                    
                    print(f"⚠️ Rate limit exceeded. Retrying in {retry_delay} seconds...")
                    time.sleep(retry_delay)
                    continue
                else:
                    # Last attempt failed
                    return {
                        'error': 'Rate limit exceeded. Please wait a few minutes and try again. Free tier has strict limits.',
                        'analysis': None,
                        'suggestion': 'Wait 1-2 minutes before trying again, or upgrade your Gemini API plan.'
                    }
            else:
                # Other errors - don't retry
                print(f"❌ Error getting Gemini analysis: {error_str}")
                return {
                    'error': f'Failed to get medical analysis: {error_str}',
                    'analysis': None
                }
    
    # Should not reach here, but just in case
    return {
        'error': 'Failed to get medical analysis after multiple attempts',
        'analysis': None
    }


def get_recommended_doctors(condition_name, limit=5):
    """
    Get recommended doctors based on the detected condition
    
    Args:
        condition_name: Name of the detected condition
        limit: Maximum number of doctors to return
    
    Returns:
        Dictionary with recommended doctors
    """
    if doctors_df is None:
        return {
            'error': 'Doctors database not loaded',
            'doctors': []
        }
    
    try:
        # Get condition code from mapping
        condition_code = CONDITION_CODE_MAP.get(condition_name)
        
        if not condition_code:
            return {
                'error': f'Unknown condition: {condition_name}',
                'doctors': []
            }
        
        print(f"🔍 Searching for doctors treating: {condition_name} ({condition_code})")
        
        # Filter doctors who treat this condition
        # The ham10000_classes_treated column contains comma-separated condition codes
        matching_doctors = doctors_df[
            doctors_df['ham10000_classes_treated'].str.contains(condition_code, case=False, na=False)
        ].copy()
        
        if len(matching_doctors) == 0:
            return {
                'error': f'No doctors found treating {condition_name}',
                'doctors': []
            }
        
        # Sort by rating (descending) and total_reviews (descending)
        # Higher rating and more reviews = better
        matching_doctors = matching_doctors.sort_values(
            by=['rating', 'total_reviews'],
            ascending=[False, False]
        )
        
        # Get top doctors
        top_doctors = matching_doctors.head(limit)
        
        # Convert to list of dictionaries
        doctors_list = []
        for _, doctor in top_doctors.iterrows():
            doctor_dict = {
                'doctor_id': str(doctor.get('doctor_id', '')),
                'doctor_name': str(doctor.get('doctor_name', 'Unknown')),
                'specialization': str(doctor.get('specialization', '')),
                'hospital_name': str(doctor.get('hospital_name', '')),
                'city': str(doctor.get('city', '')),
                'province': str(doctor.get('province', '')),
                'address': str(doctor.get('address', '')),
                'phone': str(doctor.get('phone', '')),
                'email': str(doctor.get('email', '')),
                'experience_years': int(doctor.get('experience_years', 0)),
                'consultation_fee_pkr': int(doctor.get('consultation_fee_pkr', 0)),
                'availability': str(doctor.get('availability', '')),
                'rating': float(doctor.get('rating', 0)),
                'total_reviews': int(doctor.get('total_reviews', 0)),
                'review_summary': str(doctor.get('review_summary', ''))
            }
            doctors_list.append(doctor_dict)
        
        # print(f"✓ Found {len(doctors_list)} recommended doctors for {condition_name}")
        
        return {
            'error': None,
            'doctors': doctors_list,
            'condition': condition_name,
            'condition_code': condition_code,
            'total_found': len(matching_doctors)
        }
        
    except Exception as e:
        print(f"❌ Error getting recommended doctors: {str(e)}")
        return {
            'error': f'Failed to get recommended doctors: {str(e)}',
            'doctors': []
        }


def prepare_image(image_file):
    """
    Preprocess image exactly as done during training
    """
    try:
        # Read image from uploaded file
        img = Image.open(io.BytesIO(image_file.read()))
        
        # Convert to RGB if necessary (handles RGBA, grayscale, etc.)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Resize to 224x224 (same as training)
        img = img.resize(IMG_SIZE)
        
        # Convert to numpy array
        img_array = np.array(img)
        
        # Convert to float32
        img_array = img_array.astype('float32')
        
        # Apply EfficientNet-specific preprocessing
        # This scales the image to the range expected by EfficientNet
        img_array = preprocess_input(img_array)
        
        # Add batch dimension: (224, 224, 3) -> (1, 224, 224, 3)
        img_array = np.expand_dims(img_array, axis=0)
        
        return img_array
    
    except Exception as e:
        raise ValueError(f"Image preprocessing failed: {str(e)}")


@app.route('/')
def home():
    """Serve the main HTML page"""
    return render_template('index.html')


@app.route('/predict', methods=['POST'])
def predict():
    """
    Main prediction endpoint
    Accepts: multipart/form-data with 'image' file
    Returns: JSON with predictions
    """
    # Check if model is loaded
    if model is None:
        return jsonify({
            'success': False,
            'error': 'Model not loaded. Please check server logs.'
        }), 500
    
    # Check if image was provided
    if 'image' not in request.files:
        return jsonify({
            'success': False,
            'error': 'No image file provided in request'
        }), 400
    
    file = request.files['image']
    
    # Check if filename is empty
    if file.filename == '':
        return jsonify({
            'success': False,
            'error': 'No image selected'
        }), 400
    
    try:
        # Preprocess the image
        img_array = prepare_image(file)
        
        # Make prediction
        print(f"Making prediction for image: {file.filename}")
        predictions = model.predict(img_array, verbose=0)
        
        # Extract probabilities and convert to native Python float
        pred_probs = predictions[0]
        
        # Find top prediction (highest confidence)
        top_index = int(np.argmax(pred_probs))
        confidence_value = float(pred_probs[top_index])
        percentage_value = float(confidence_value * 100)
        
        top_prediction = {
            'label': str(CLASS_NAMES[top_index]),
            'confidence': confidence_value,
            'percentage': round(percentage_value, 2)
        }
        
        print(f"✓ Prediction complete: {top_prediction['label']} ({top_prediction['percentage']}%)")
        
        # Get medical analysis from Gemini
        medical_analysis = get_medical_analysis(
            cancer_name=top_prediction['label'],
            confidence_score=confidence_value,
            model_accuracy=MODEL_ACCURACY
        )
        
        # Get recommended doctors based on condition
        recommended_doctors = get_recommended_doctors(
            condition_name=top_prediction['label'],
            limit=5
        )
        
        # Prepare response
        response_data = {
            'success': True,
            'top_prediction': {
                'label': str(top_prediction['label']),
                'confidence': float(top_prediction['confidence']),
                'percentage': float(top_prediction['percentage'])
            },
            'model_accuracy': MODEL_ACCURACY,
            'medical_analysis': medical_analysis,
            'recommended_doctors': recommended_doctors
        }
        
        return jsonify(response_data)
    
    except ValueError as ve:
        return jsonify({
            'success': False,
            'error': str(ve)
        }), 400
    
    except Exception as e:
        print(f"❌ Prediction error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Prediction failed: {str(e)}'
        }), 500


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'model_loaded': model is not None,
        'model_path': MODEL_PATH,
        'num_classes': len(CLASS_NAMES)
    })


@app.route('/classes', methods=['GET'])
def get_classes():
    """Return all class names"""
    return jsonify({
        'classes': CLASS_NAMES,
        'count': len(CLASS_NAMES)
    })


@app.route('/chat', methods=['POST'])
def chat():
    """
    Chat endpoint for asking questions about the medical analysis
    Accepts: JSON with 'question' and 'medical_analysis' context
    Returns: JSON with AI response
    """
    if not GEMINI_API_KEY:
        return jsonify({
            'success': False,
            'error': 'Gemini API key not configured'
        }), 500
    
    try:
        data = request.get_json()
        question = data.get('question', '').strip()
        medical_analysis = data.get('medical_analysis', '')
        condition_name = data.get('condition_name', '')
        confidence_score = data.get('confidence_score', 0)
        
        if not question:
            return jsonify({
                'success': False,
                'error': 'No question provided'
            }), 400
        
        if not medical_analysis:
            return jsonify({
                'success': False,
                'error': 'No medical analysis context provided. Please upload an image first.'
            }), 400
        
        # Create prompt for chatbot
        prompt = f"""You are a helpful medical AI assistant helping a user understand their skin condition analysis.

CONTEXT - Medical Analysis Report:
Condition: {condition_name}
Confidence: {confidence_score:.2%}
Model Accuracy: {MODEL_ACCURACY}%

Medical Analysis:
{medical_analysis}

USER QUESTION: {question}

INSTRUCTIONS:
- Answer the user's question based on the medical analysis provided above
- Be clear, concise, and helpful
- Use simple language that's easy to understand
- If the question is not related to the medical analysis, politely redirect them
- Always remind them that this is AI-assisted information and they should consult a qualified dermatologist
- If you don't know something, say so honestly
- Keep responses under 300 words unless more detail is specifically requested

Provide a helpful, accurate answer:"""

        # Initialize the model
        gemini_model = genai.GenerativeModel(GEMINI_MODEL)
        
        # Generate response with retry logic
        max_retries = 2
        for attempt in range(max_retries):
            try:
                print(f"💬 Chat question: {question[:50]}... (Attempt {attempt + 1}/{max_retries})")
                response = gemini_model.generate_content(prompt)
                answer = response.text
                
                print(f"✓ Chat response generated")
                
                return jsonify({
                    'success': True,
                    'answer': answer
                })
                
            except Exception as e:
                error_str = str(e)
                
                # Check if it's a rate limit error
                if ('429' in error_str or 'quota' in error_str.lower() or 'rate limit' in error_str.lower()) and attempt < max_retries - 1:
                    retry_delay = 10  # Shorter delay for chat
                    print(f"⚠️ Rate limit. Retrying in {retry_delay} seconds...")
                    time.sleep(retry_delay)
                    continue
                else:
                    raise e
        
    except Exception as e:
        print(f"❌ Chat error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Failed to get answer: {str(e)}'
        }), 500


@app.route('/download_report', methods=['POST'])
def download_report():
    """
    Generate and return a report in PDF or DOCX format with patient information.
    Expects JSON body with keys: patientName, patientAge, patientSex, 
    detectedCondition, confidencePercent, modelAccuracy,
    medicalAnalysis, doctors (list), doctorsMeta (optional). 
    Query param: format=pdf|docx
    """
    fmt = request.args.get('format', 'pdf').lower()
    data = request.get_json(silent=True) or {}

    # Extract patient information
    patient_name = data.get('patientName', 'Unknown Patient')
    patient_age = data.get('patientAge', 'N/A')
    patient_sex = data.get('patientSex', 'N/A')

    if fmt == 'pdf':
        if not REPORTLAB_AVAILABLE:
            return jsonify({'success': False, 'error': 'PDF backend not installed (reportlab).'}), 500
        try:
            buffer = io.BytesIO()
            page_width, page_height = A4
            c = canvas.Canvas(buffer, pagesize=A4)

            # Letterhead bar - Medical Blue
            c.setFillColorRGB(8/255.0, 145/255.0, 178/255.0)
            c.rect(0, page_height - 80, page_width, 80, stroke=0, fill=1)

            # Title
            c.setFillColorRGB(1, 1, 1)
            c.setFont("Helvetica-Bold", 18)
            c.drawString(40, page_height - 35, "Skin Lesion Detection Report")
            c.setFont("Helvetica", 9)
            c.drawRightString(page_width - 40, page_height - 35, "AI-Powered Dermatological Analysis")
            c.drawRightString(page_width - 40, page_height - 50, "Created by Syed Atif Shah, Daim Ahmad")
            c.drawRightString(page_width - 40, page_height - 65, f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")

            y = page_height - 100
            c.setFillColorRGB(0, 0, 0)

            # Patient Information Section - Highlighted Box
            c.setFillColorRGB(240/255.0, 249/255.0, 255/255.0)  # Light blue background
            c.rect(40, y - 80, page_width - 80, 75, stroke=1, fill=1)
            c.setFillColorRGB(0, 0, 0)
            
            c.setFont("Helvetica-Bold", 13)
            c.drawString(50, y - 15, "Patient Information")
            c.setFont("Helvetica", 11)
            c.drawString(50, y - 35, f"Name: {patient_name}")
            c.drawString(50, y - 50, f"Age: {patient_age} years")
            c.drawString(50, y - 65, f"Sex: {patient_sex}")
            
            y -= 100

            # Detection Results
            c.setFont("Helvetica-Bold", 13)
            c.drawString(40, y, "Detection Results")
            y -= 20
            c.setFont("Helvetica", 11)
            lines = [
                f"Detected Condition: {data.get('detectedCondition','')}",
                f"Confidence Score: {data.get('confidencePercent','')}%",
                f"Model Accuracy: {data.get('modelAccuracy','')}%",
            ]
            for line in lines:
                c.drawString(40, y, line)
                y -= 16
            y -= 10

            # Analysis text wrapped
            c.setFont("Helvetica-Bold", 13)
            c.drawString(40, y, "Specialized Medical LLM Analysis")
            y -= 20
            c.setFont("Helvetica", 10)
            from reportlab.lib.utils import simpleSplit
            analysis_text = data.get('medicalAnalysis') or 'Not available.'
            wrap = simpleSplit(analysis_text, "Helvetica", 10, page_width - 80)
            for line in wrap:
                if y < 60:
                    c.showPage()
                    # Small header on new page
                    c.setFillColorRGB(8/255.0, 145/255.0, 178/255.0)
                    c.rect(0, page_height - 30, page_width, 30, stroke=0, fill=1)
                    c.setFillColorRGB(1, 1, 1)
                    c.setFont("Helvetica-Bold", 12)
                    c.drawString(40, page_height - 18, f"Patient: {patient_name} (continued)")
                    y = page_height - 50
                    c.setFillColorRGB(0, 0, 0)
                    c.setFont("Helvetica", 10)
                c.drawString(40, y, line)
                y -= 14
            y -= 10

            # Doctors Section
            if y < 100:
                c.showPage()
                c.setFillColorRGB(8/255.0, 145/255.0, 178/255.0)
                c.rect(0, page_height - 30, page_width, 30, stroke=0, fill=1)
                c.setFillColorRGB(1, 1, 1)
                c.setFont("Helvetica-Bold", 12)
                c.drawString(40, page_height - 18, f"Patient: {patient_name} (continued)")
                y = page_height - 50
                c.setFillColorRGB(0, 0, 0)
            
            c.setFont("Helvetica-Bold", 13)
            c.drawString(40, y, "Recommended Specialists")
            y -= 20
            c.setFont("Helvetica", 10)
            doctors = data.get('doctors') or []
            if not doctors:
                c.drawString(40, y, "No specialists available for this condition.")
                y -= 18
            else:
                for i, d in enumerate(doctors, 1):
                    if y < 150:
                        c.showPage()
                        c.setFillColorRGB(8/255.0, 145/255.0, 178/255.0)
                        c.rect(0, page_height - 30, page_width, 30, stroke=0, fill=1)
                        c.setFillColorRGB(1, 1, 1)
                        c.setFont("Helvetica-Bold", 12)
                        c.drawString(40, page_height - 18, f"Patient: {patient_name} (continued)")
                        y = page_height - 50
                        c.setFillColorRGB(0, 0, 0)
                        c.setFont("Helvetica", 10)
                    
                    # Doctor name in bold
                    c.setFont("Helvetica-Bold", 11)
                    c.drawString(40, y, f"{i}. {d.get('doctor_name','Doctor')}")
                    y -= 16
                    c.setFont("Helvetica", 10)
                    
                    block = []
                    if d.get('specialization'): block.append(f"   Specialization: {d.get('specialization')}")
                    if d.get('hospital_name'): block.append(f"   Hospital: {d.get('hospital_name')}")
                    loc = ", ".join([v for v in [d.get('address'), d.get('city'), d.get('province')] if v])
                    if loc: block.append(f"   Location: {loc}")
                    rating = d.get('rating')
                    if isinstance(rating, (int, float)): block.append(f"   Rating: {rating:.1f} ({d.get('total_reviews',0)} reviews)")
                    if d.get('experience_years'): block.append(f"   Experience: {int(d.get('experience_years'))} years")
                    if d.get('consultation_fee_pkr'): block.append(f"   Consultation Fee: PKR {int(d.get('consultation_fee_pkr'))}")
                    if d.get('availability'): block.append(f"   Availability: {d.get('availability')}")
                    if d.get('phone') and d.get('phone') != 'Not available': block.append(f"   Phone: {d.get('phone')}")
                    if d.get('email') and d.get('email') != 'Not available': block.append(f"   Email: {d.get('email')}")

                    for line in block:
                        if y < 60:
                            c.showPage()
                            c.setFillColorRGB(8/255.0, 145/255.0, 178/255.0)
                            c.rect(0, page_height - 30, page_width, 30, stroke=0, fill=1)
                            c.setFillColorRGB(1, 1, 1)
                            c.setFont("Helvetica-Bold", 12)
                            c.drawString(40, page_height - 18, f"Patient: {patient_name} (continued)")
                            y = page_height - 50
                            c.setFillColorRGB(0, 0, 0)
                            c.setFont("Helvetica", 10)
                        c.drawString(40, y, line)
                        y -= 14
                    y -= 8

            # Disclaimer
            if y < 100:
                c.showPage()
                y = page_height - 60
            c.setStrokeColorRGB(0.7, 0.7, 0.7)
            c.line(40, y, page_width - 40, y)
            y -= 16
            c.setFont("Helvetica-Oblique", 9)
            c.setFillColorRGB(0.3, 0.3, 0.3)
            disc = "Medical Disclaimer: This AI-assisted report is for educational purposes only and does not constitute professional medical advice, diagnosis, or treatment. Always consult a qualified dermatologist for proper medical evaluation and care."
            wrap = simpleSplit(disc, "Helvetica-Oblique", 9, page_width - 80)
            for line in wrap:
                c.drawString(40, y, line)
                y -= 12

            c.save()
            buffer.seek(0)
            
            # Create filename with patient name
            safe_name = re.sub(r'[^a-zA-Z0-9_]', '_', patient_name)
            safe_condition = re.sub(r'[^a-zA-Z0-9_]', '_', data.get('detectedCondition', 'condition'))
            timestamp = int(time.time())
            filename = f"Medical_Report_{safe_name}_{safe_condition}_{timestamp}.pdf"
            
            return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')
        except Exception as e:
            print(f"❌ PDF generation error: {str(e)}")
            return jsonify({'success': False, 'error': f'PDF generation failed: {str(e)}'}), 500

    # DOCX Format
    if fmt == 'docx':
        if not DOCX_AVAILABLE:
            return jsonify({'success': False, 'error': 'DOCX backend not installed (python-docx).'}), 500
        try:
            doc = Document()
            
            # Title with styling
            doc.add_heading('Skin Lesion Detection Report', level=0)
            
            subtitle = doc.add_paragraph('AI-Powered Dermatological Analysis')
            subtitle_format = subtitle.runs[0]
            subtitle_format.font.italic = True
            
            # Authors and date
            doc.add_paragraph('Created by Syed Atif Shah, Daim Ahmad')
            doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            
            doc.add_paragraph()  # Spacer

            # Patient Information Section
            doc.add_heading('Patient Information', level=1)
            patient_table = doc.add_table(rows=3, cols=2)
            patient_table.style = 'Light Grid Accent 1'
            
            patient_table.rows[0].cells[0].text = 'Patient Name:'
            patient_table.rows[0].cells[1].text = str(patient_name)
            patient_table.rows[1].cells[0].text = 'Age:'
            patient_table.rows[1].cells[1].text = f"{patient_age} years"
            patient_table.rows[2].cells[0].text = 'Sex:'
            patient_table.rows[2].cells[1].text = str(patient_sex)
            
            # Make first column bold
            for row in patient_table.rows:
                row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()  # Spacer

            # Detection Results
            doc.add_heading('Detection Results', level=1)
            doc.add_paragraph(f"Detected Condition: {data.get('detectedCondition','')}")
            doc.add_paragraph(f"Confidence Score: {data.get('confidencePercent','')}%")
            doc.add_paragraph(f"Model Accuracy: {data.get('modelAccuracy','')}%")
            
            doc.add_paragraph()  # Spacer

            # Medical Analysis
            doc.add_heading('Specialized Medical LLM Analysis', level=1)
            doc.add_paragraph(data.get('medicalAnalysis') or 'Not available.')

            doc.add_paragraph()  # Spacer

            # Recommended Doctors
            doc.add_heading('Recommended Specialists', level=1)
            doctors = data.get('doctors') or []
            if not doctors:
                doc.add_paragraph('No specialists available for this condition.')
            else:
                for i, d in enumerate(doctors, 1):
                    # Doctor name as heading
                    p = doc.add_paragraph()
                    p.add_run(f"{i}. {d.get('doctor_name','Doctor')}\n").bold = True
                    
                    # Doctor details
                    if d.get('specialization'): 
                        doc.add_paragraph(f"Specialization: {d.get('specialization')}")
                    if d.get('hospital_name'): 
                        doc.add_paragraph(f"Hospital: {d.get('hospital_name')}")
                    
                    loc = ", ".join([v for v in [d.get('address'), d.get('city'), d.get('province')] if v])
                    if loc: 
                        doc.add_paragraph(f"Location: {loc}")
                    
                    rating = d.get('rating')
                    if isinstance(rating, (int, float)): 
                        doc.add_paragraph(f"Rating: {rating:.1f} ({d.get('total_reviews',0)} reviews)")
                    
                    if d.get('experience_years'): 
                        doc.add_paragraph(f"Experience: {int(d.get('experience_years'))} years")
                    if d.get('consultation_fee_pkr'): 
                        doc.add_paragraph(f"Consultation Fee: PKR {int(d.get('consultation_fee_pkr'))}")
                    if d.get('availability'): 
                        doc.add_paragraph(f"Availability: {d.get('availability')}")
                    if d.get('phone') and d.get('phone') != 'Not available': 
                        doc.add_paragraph(f"Phone: {d.get('phone')}")
                    if d.get('email') and d.get('email') != 'Not available': 
                        doc.add_paragraph(f"Email: {d.get('email')}")
                    
                    doc.add_paragraph()  # Spacer between doctors

            # Disclaimer
            doc.add_paragraph()
            disclaimer = doc.add_paragraph()
            disclaimer.add_run('Medical Disclaimer: ').bold = True
            disclaimer.add_run('This AI-assisted report is for educational purposes only and does not constitute professional medical advice, diagnosis, or treatment. Always consult a qualified dermatologist for proper medical evaluation and care.')

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Create filename with patient name
            safe_name = re.sub(r'[^a-zA-Z0-9_]', '_', patient_name)
            safe_condition = re.sub(r'[^a-zA-Z0-9_]', '_', data.get('detectedCondition', 'condition'))
            timestamp = int(time.time())
            filename = f"Medical_Report_{safe_name}_{safe_condition}_{timestamp}.docx"
            
            return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        except Exception as e:
            print(f"❌ DOCX generation error: {str(e)}")
            return jsonify({'success': False, 'error': f'DOCX generation failed: {str(e)}'}), 500

    return jsonify({'success': False, 'error': 'Unsupported format. Use format=pdf or format=docx'}), 400


if __name__ == '__main__':
    print("\n" + "="*60)
    print("🚀 STARTING FLASK SERVER")
    print("="*60)
    print(f"📍 URL: http://localhost:5000")
    print(f"📍 Health Check: http://localhost:5000/health")
    print(f"📍 Prediction API: http://localhost:5000/predict")
    print(f"📍 Chat API: http://localhost:5000/chat")
    print("="*60 + "\n")
    
    # Run the Flask app
    app.run(
        debug=True,
        host='0.0.0.0',
        port=5000,
        threaded=True
    )